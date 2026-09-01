from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).with_name('P1.docx')
ACCENT='1D4ED8'; DARK='0F172A'; MUTED='64748B'; LIGHT='EFF6FF'; BORDER='CBD5E1'; GREEN='15803D'; FIELD='F8FAFC'

def shade(cell, fill):
    pr=cell._tc.get_or_add_tcPr(); shd=pr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); pr.append(shd)
    shd.set(qn('w:fill'),fill)

def borders(cell, **edges):
    pr=cell._tc.get_or_add_tcPr(); b=pr.first_child_found_in('w:tcBorders')
    if b is None: b=OxmlElement('w:tcBorders'); pr.append(b)
    for edge,data in edges.items():
        el=b.find(qn('w:'+edge))
        if el is None: el=OxmlElement('w:'+edge); b.append(el)
        for k,v in data.items(): el.set(qn('w:'+k),str(v))

def margins(cell, top=70, start=90, bottom=70, end=90):
    pr=cell._tc.get_or_add_tcPr(); m=pr.first_child_found_in('w:tcMar')
    if m is None: m=OxmlElement('w:tcMar'); pr.append(m)
    for name,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        el=m.find(qn('w:'+name))
        if el is None: el=OxmlElement('w:'+name); m.append(el)
        el.set(qn('w:w'),str(val)); el.set(qn('w:type'),'dxa')

def cant_split(row):
    pr=row._tr.get_or_add_trPr()
    if pr.find(qn('w:cantSplit')) is None: pr.append(OxmlElement('w:cantSplit'))

def keep_next(p):
    pr=p._p.get_or_add_pPr()
    if pr.find(qn('w:keepNext')) is None: pr.append(OxmlElement('w:keepNext'))

def widths(table, values):
    table.autofit=False; pr=table._tbl.tblPr
    lay=pr.find(qn('w:tblLayout'))
    if lay is None: lay=OxmlElement('w:tblLayout'); pr.append(lay)
    lay.set(qn('w:type'),'fixed')
    tw=pr.find(qn('w:tblW'))
    if tw is None: tw=OxmlElement('w:tblW'); pr.append(tw)
    tw.set(qn('w:type'),'dxa'); tw.set(qn('w:w'),str(Cm(sum(values)).twips))
    for i,w in enumerate(values):
        table.columns[i].width=Cm(w)
        for row in table.rows: row.cells[i].width=Cm(w); cant_split(row)

def set_row_height(row, cm):
    row.height=Cm(cm); row.height_rule=WD_ROW_HEIGHT_RULE.AT_LEAST

def page_field(p):
    r=p.add_run('Seite '); r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(MUTED)
    a=OxmlElement('w:fldChar'); a.set(qn('w:fldCharType'),'begin')
    t=OxmlElement('w:instrText'); t.set(qn('xml:space'),'preserve'); t.text=' PAGE '
    z=OxmlElement('w:fldChar'); z.set(qn('w:fldCharType'),'end')
    r._r.extend([a,t,z])

def empty_field_cell(cell, height=.68, fill=FIELD, border_color=BORDER):
    shade(cell, fill); margins(cell,60,85,60,85)
    edge={'val':'single','sz':'5','color':border_color}
    borders(cell,top=edge,bottom=edge,left=edge,right=edge)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p=cell.paragraphs[0]
    p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
    p.add_run('')
    set_row_height(cell._tc.getparent().getparent(), height)

def answer_box(doc, height=1.0):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; widths(t,[17.2])
    c=t.cell(0,0); empty_field_cell(c,height=height)
    t.rows[0].height=Cm(height); t.rows[0].height_rule=WD_ROW_HEIGHT_RULE.AT_LEAST
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    return t

def labeled_answer_row(doc, label, width_label=5.2, height=.7):
    t=doc.add_table(rows=1,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; widths(t,[width_label,17.2-width_label])
    c0,c1=t.rows[0].cells
    margins(c0,55,40,55,60); borders(c0,bottom={'val':'single','sz':'4','color':'E2E8F0'})
    p=c0.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(label); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(DARK)
    empty_field_cell(c1,height=height)
    t.rows[0].height=Cm(height); t.rows[0].height_rule=WD_ROW_HEIGHT_RULE.AT_LEAST
    return t

doc=Document(); sec=doc.sections[0]
sec.page_width=Cm(21); sec.page_height=Cm(29.7); sec.top_margin=Cm(1.15); sec.bottom_margin=Cm(1.05); sec.left_margin=Cm(1.55); sec.right_margin=Cm(1.55)
styles=doc.styles; styles['Normal'].font.name='Arial'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial'); styles['Normal'].font.size=Pt(9.5); styles['Normal'].paragraph_format.space_after=Pt(3)
for name in ['Heading 1','Heading 2','Heading 3']:
    styles[name].font.name='Arial'; styles[name]._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial')
styles['Heading 1'].font.size=Pt(14); styles['Heading 1'].font.bold=True; styles['Heading 1'].font.color.rgb=RGBColor.from_string(DARK); styles['Heading 1'].paragraph_format.space_before=Pt(7); styles['Heading 1'].paragraph_format.space_after=Pt(3)
styles['Heading 2'].font.size=Pt(10.5); styles['Heading 2'].font.bold=True; styles['Heading 2'].font.color.rgb=RGBColor.from_string(ACCENT); styles['Heading 2'].paragraph_format.space_before=Pt(4); styles['Heading 2'].paragraph_format.space_after=Pt(2)

p=sec.footer.paragraphs[0]; p.text='P1 · Übungstest Tastenkombinationen · '; p.runs[0].font.name='Arial'; p.runs[0].font.size=Pt(8); p.runs[0].font.color.rgb=RGBColor.from_string(MUTED); page_field(p)

# Header band
band=doc.add_table(rows=1,cols=1); band.alignment=WD_TABLE_ALIGNMENT.CENTER; widths(band,[17.2]); c=band.cell(0,0); shade(c,DARK); margins(c,150,200,145,200)
p=c.paragraphs[0]; r=p.add_run('P1  |  ÜBUNGSTEST TASTENKOMBINATIONEN'); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=RGBColor(255,255,255); r.font.name='Arial'
p=c.add_paragraph(); p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(0); r=p.add_run('Tastenkombinationen A1-A7  ·  ca. 35-40 Minuten  ·  ohne Kursseiten oder Merkblatt'); r.font.size=Pt(8.8); r.font.color.rgb=RGBColor(203,213,225); r.font.name='Arial'

# Student metadata: actual empty cells, no underscores to delete
meta=doc.add_table(rows=2,cols=3); meta.alignment=WD_TABLE_ALIGNMENT.CENTER; widths(meta,[8.0,4.0,5.2])
for i,label in enumerate(['VORNAME / NAME','KLASSE','DATUM']):
    c=meta.cell(0,i); margins(c,45,70,15,70); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(label); r.font.size=Pt(7.3); r.bold=True; r.font.color.rgb=RGBColor.from_string(MUTED)
    empty_field_cell(meta.cell(1,i),height=.66)
meta.rows[1].height=Cm(.66)

score=doc.add_table(rows=1,cols=3); score.alignment=WD_TABLE_ALIGNMENT.CENTER; widths(score,[8.0,4.0,5.2])
for i,(lab,val) in enumerate([('ERREICHTE PUNKTE',''),('GESAMTPUNKTE','41'),('NOTE','')]):
    c=score.cell(0,i); shade(c,LIGHT if i!=2 else 'ECFDF5'); margins(c,75,90,75,90)
    edge={'val':'single','sz':'5','color':'BFDBFE'}; borders(c,top=edge,bottom=edge,left=edge,right=edge)
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(lab+'\n'); r.font.size=Pt(7.2); r.bold=True; r.font.color.rgb=RGBColor.from_string(MUTED)
    r=p.add_run(val); r.font.size=Pt(13); r.bold=True; r.font.color.rgb=RGBColor.from_string(GREEN if i==2 else DARK)

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(2); r=p.add_run('NOTENSCHLÜSSEL'); r.bold=True; r.font.size=Pt(7.6); r.font.color.rgb=RGBColor.from_string(MUTED)
t=doc.add_table(rows=2,cols=6); t.alignment=WD_TABLE_ALIGNMENT.CENTER; widths(t,[17.2/6]*6)
for j,(pts,grade) in enumerate(zip(['41 P','33 P','25 P','17 P','9 P','0 P'],['6.0','5.0','4.0','3.0','2.0','1.0'])):
    for i,text in enumerate([pts,grade]):
        c=t.cell(i,j); margins(c,45,45,45,45); edge={'val':'single','sz':'4','color':BORDER}; borders(c,top=edge,bottom=edge,left=edge,right=edge)
        if i==0: shade(c,'F8FAFC')
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); r=p.add_run(text); r.font.size=Pt(8.1); r.bold=(i==1); r.font.color.rgb=RGBColor.from_string(DARK if i else MUTED)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(3); r=p.add_run('Linear: Note = 1 + 5 × (erreichte Punkte / 41)  ·  60 % = Note 4.0'); r.font.size=Pt(7.1); r.font.color.rgb=RGBColor.from_string(MUTED)

def heading(n,title,pts,break_before=False):
    p=doc.add_paragraph(style='Heading 1'); p.paragraph_format.page_break_before=break_before; keep_next(p); p.add_run(f'{n}. {title}'); r=p.add_run(f'   {pts} P'); r.font.size=Pt(9.3); r.font.color.rgb=RGBColor.from_string(ACCENT)

def instruction(text):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); keep_next(p); r=p.add_run(text); r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(MUTED)

def mc(title,opts):
    p=doc.add_paragraph(style='Heading 2'); p.add_run(title); keep_next(p)
    for opt in opts:
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(.12); p.paragraph_format.space_after=Pt(.8); p.add_run('☐  '+opt).font.size=Pt(8.7)

def scenario(title,text,opts):
    p=doc.add_paragraph(style='Heading 2'); p.add_run(title); keep_next(p)
    p=doc.add_paragraph(text); p.paragraph_format.space_after=Pt(1.2); p.runs[0].font.size=Pt(8.8)
    for opt in opts:
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(.15); p.paragraph_format.space_after=Pt(.7); p.add_run('☐  '+opt).font.size=Pt(8.8)

# PAGE 1
heading(1,'Was macht dieses Tastenkürzel?',4); instruction('Kreuze jeweils die richtige Aussage an. Genau eine Aussage ist korrekt.')
for title,opts in [
('Ctrl + C',['Kopiert markierte Inhalte.','Fügt kopierte Inhalte ein.','Schliesst die aktuelle Anwendung.']),
('Ctrl + V',['Speichert das Dokument.','Fügt kopierte oder ausgeschnittene Inhalte ein.','Macht den letzten Schritt rückgängig.']),
('Ctrl + X',['Schneidet markierte Inhalte aus.','Sucht nach einem Begriff.','Öffnet eine Datei.']),
('Ctrl + Z',['Macht den letzten Schritt rückgängig.','Markiert alles.','Druckt das Dokument.'])]: mc(title,opts)

# PAGE 2
heading(2,'Tastenkombinationen im Alltag',6); instruction('Wähle jeweils die sinnvollste Tastenkombination.')
for row in [
('1. Dokument sichern','Elena arbeitet seit längerer Zeit an einem wichtigen Dokument. Bevor sie weiterarbeitet, möchte sie den aktuellen Stand speichern.',['Ctrl + S','Ctrl + P','Ctrl + O']),
('2. Alles formatieren','Jan möchte den gesamten Text seines Dokuments markieren, um die Schriftart für alles gleichzeitig zu ändern.',['Ctrl + A','Ctrl + F','Ctrl + C']),
('3. Begriff finden','Amir hat ein sechsseitiges Dokument geöffnet. Er sucht darin nach dem Begriff „Datenschutz“.',['Ctrl + F','Ctrl + H','Ctrl + L']),
('4. Direkt ans Ende','Sophie arbeitet an einem sehr langen Bericht. Sie befindet sich weit oben und möchte direkt ganz ans Ende des Dokuments springen.',['Ctrl + End','Ctrl + Home','Ctrl + ↓']),
('5. Ohne fremde Formatierung','Tim kopiert Text von einer Webseite in ein Dokument. Die fremde Schriftart und Formatierung soll möglichst nicht übernommen werden.',['Ctrl + Shift + V','Ctrl + V','Ctrl + Shift + C']),
('6. Doch wiederherstellen','Luca hat einen Schritt rückgängig gemacht und merkt, dass er ihn doch behalten wollte.',['Ctrl + Y','Ctrl + Z','Ctrl + S'])]: scenario(*row)

# PAGE 3
heading(3,'Sonderzeichen mit AltGr',8); instruction('Ergänze jeweils die zweite Taste auf einer Schweizer Tastatur. Das Zielzeichen bleibt sichtbar; geprüft wird nur die zweite Taste zu AltGr.')
t=doc.add_table(rows=0,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; widths(t,[6.2,11.0])
for i,char in enumerate(['@','#','€','|','\\','[','{','°'],1):
    cells=t.add_row().cells
    margins(cells[0],50,80,50,80); borders(cells[0],bottom={'val':'single','sz':'4','color':'E2E8F0'})
    p=cells[0].paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.add_run(f'{i}.  {char}').font.size=Pt(9)
    empty_field_cell(cells[1],height=.57)
    t.rows[-1].height=Cm(.57); t.rows[-1].height_rule=WD_ROW_HEIGHT_RULE.AT_LEAST

heading(4,'Programme & Browser',4); instruction('Wähle die passendste Tastenkombination.')
for row in [
('1. Neues Word-Dokument','Du möchtest in Word ein neues Dokument erstellen.',['Ctrl + N','Ctrl + T','Ctrl + W']),
('2. Neuer Browser-Tab','Du möchtest im Browser einen neuen Tab öffnen.',['Ctrl + N','Ctrl + T','Ctrl + L']),
('3. Webseite neu laden','Du möchtest die aktuelle Webseite neu laden.',['F5','Ctrl + F','Win + D']),
('4. Adresse markieren','Du möchtest sofort die gesamte Adresse im Browser markieren.',['Ctrl + L','Ctrl + F','Ctrl + T'])]: scenario(*row)

# PAGE 4 - intentionally compact: sections 5-8 + bonus share one page
heading(5,'Windows & Arbeitsalltag',5,True); instruction('Schreibe die passende Tastenkombination direkt ins leere Feld.')
for i,text in enumerate(['Desktop anzeigen','Datei-Explorer öffnen','Einen Bildschirmausschnitt aufnehmen','Task-Manager direkt öffnen','Verlauf der Zwischenablage öffnen'],1):
    labeled_answer_row(doc,f'{i}. {text}',width_label=11.7,height=.53)

heading(6,'Kurz erklären',6); instruction('Beschreibe kurz die Wirkung. Stichwort oder Satz genügt; bewertet wird der Inhalt. Je 1 Punkt.')
for i,key in enumerate(['Ctrl + H in Word','Ctrl + P','Ctrl + O','Win + ←','Win + ↑','Win + ↓ (Fenster maximiert)'],1):
    labeled_answer_row(doc,f'{i}. {key}',width_label=5.2,height=.53)

heading(7,'Nicht verwechseln',4); instruction('Erkläre jeweils kurz den Unterschied. Pro Teilaufgabe sind 2 Punkte möglich.')
p=doc.add_paragraph(style='Heading 2'); p.add_run('A) Ctrl + Tab und Alt + Tab'); keep_next(p)
p=doc.add_paragraph('Wann verwendest du Ctrl + Tab und wann Alt + Tab?'); p.paragraph_format.space_after=Pt(2); p.runs[0].font.size=Pt(8.7); answer_box(doc,.78)
p=doc.add_paragraph(style='Heading 2'); p.add_run('B) Ctrl + W und Alt + F4'); keep_next(p)
p=doc.add_paragraph('Was schliesst Ctrl + W typischerweise und was schliesst Alt + F4?'); p.paragraph_format.space_after=Pt(2); p.runs[0].font.size=Pt(8.7); answer_box(doc,.78)

heading(8,'Transfer',4); instruction('Nenne die sinnvollste Tastenkombination und begründe kurz.')
for title,text in [
('A) Zwei Programme nebeneinander - 2 P','Das aktive Browserfenster soll ohne Maus die rechte Bildschirmhälfte einnehmen, damit PowerPoint daneben sichtbar bleiben kann.'),
('B) Arbeitsplatz verlassen - 2 P','Du gehst kurz aus dem Raum. Programme sollen geöffnet bleiben, aber niemand soll den Computer benutzen können.')]:
    p=doc.add_paragraph(style='Heading 2'); p.add_run(title); keep_next(p)
    p=doc.add_paragraph(text); p.paragraph_format.space_after=Pt(1); p.runs[0].font.size=Pt(8.6)
    labeled_answer_row(doc,'Tastenkombination',width_label=4.2,height=.48)
    labeled_answer_row(doc,'Begründung',width_label=4.2,height=.55)

p=doc.add_paragraph(style='Heading 1'); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(1); p.add_run('Bonus - max. +2 P')
p=doc.add_paragraph('Genau unterscheiden: je +1 P.'); p.paragraph_format.space_after=Pt(2); p.runs[0].font.size=Pt(8.4); p.runs[0].font.color.rgb=RGBColor.from_string(MUTED)
bonus=doc.add_table(rows=0,cols=2); bonus.alignment=WD_TABLE_ALIGNMENT.CENTER; widths(bonus,[12.5,4.7])
items=[
('Bonus 1 · Im Browser bleiben (+1 P)','Zum vorherigen Browser-Tab wechseln, ohne zu Word zu wechseln.'),
('Bonus 2 · Fenster neu anordnen (+1 P)','Maximiertes Fenster zuerst normal gross machen, danach links andocken. Zwei Kürzel in Reihenfolge.')]
for title,text in items:
    cells=bonus.add_row().cells
    left,right=cells
    margins(left,70,90,70,90); margins(right,55,75,55,75)
    edge={'val':'single','sz':'5','color':BORDER}
    borders(left,top=edge,bottom=edge,left=edge,right=edge); borders(right,top=edge,bottom=edge,left=edge,right=edge)
    p=left.paragraphs[0]; p.paragraph_format.space_after=Pt(1); r=p.add_run(title); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(ACCENT)
    p=left.add_paragraph(text); p.paragraph_format.space_after=Pt(0); p.runs[0].font.size=Pt(7.9)
    shade(right,'F8FAFC')
    p=right.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run('Antwort'); r.bold=True; r.font.size=Pt(7.4); r.font.color.rgb=RGBColor.from_string(MUTED)
    p=right.add_paragraph(''); p.paragraph_format.space_after=Pt(0)
    bonus.rows[-1].height=Cm(1.08); bonus.rows[-1].height_rule=WD_ROW_HEIGHT_RULE.AT_LEAST
p=doc.add_paragraph('Bonuspunkte zusätzlich zu 41 P · Maximalnote 6.0'); p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(0); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].font.size=Pt(7.3); p.runs[0].font.color.rgb=RGBColor.from_string(MUTED)

doc.core_properties.title='P1 - Übungstest Tastenkombinationen'; doc.core_properties.subject='Tastenkombinationen A1-A7'; doc.core_properties.author=''; doc.core_properties.keywords='Tastenkombinationen, Übungstest, Informatik'
OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT); print(OUT)
