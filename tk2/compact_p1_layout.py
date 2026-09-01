from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

PATH = Path(__file__).with_name('P1.docx')
ACCENT='1D4ED8'; DARK='0F172A'; MUTED='64748B'; BORDER='CBD5E1'


def margins(cell, top=70, start=100, bottom=70, end=100):
    pr=cell._tc.get_or_add_tcPr(); m=pr.first_child_found_in('w:tcMar')
    if m is None: m=OxmlElement('w:tcMar'); pr.append(m)
    for name,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        el=m.find(qn('w:'+name))
        if el is None: el=OxmlElement('w:'+name); m.append(el)
        el.set(qn('w:w'),str(val)); el.set(qn('w:type'),'dxa')


def border(cell, color=BORDER, size='4'):
    pr=cell._tc.get_or_add_tcPr(); b=pr.first_child_found_in('w:tcBorders')
    if b is None: b=OxmlElement('w:tcBorders'); pr.append(b)
    for edge in ('top','left','bottom','right'):
        el=b.find(qn('w:'+edge))
        if el is None: el=OxmlElement('w:'+edge); b.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),size); el.set(qn('w:color'),color)


def shade(cell, fill='F8FAFC'):
    pr=cell._tc.get_or_add_tcPr(); el=pr.find(qn('w:shd'))
    if el is None: el=OxmlElement('w:shd'); pr.append(el)
    el.set(qn('w:fill'),fill)


def cant_split(row):
    pr=row._tr.get_or_add_trPr()
    if pr.find(qn('w:cantSplit')) is None: pr.append(OxmlElement('w:cantSplit'))


def fixed(table, widths):
    table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    pr=table._tbl.tblPr; lay=pr.find(qn('w:tblLayout'))
    if lay is None: lay=OxmlElement('w:tblLayout'); pr.append(lay)
    lay.set(qn('w:type'),'fixed')
    for i,w in enumerate(widths):
        table.columns[i].width=Cm(w)
        for row in table.rows:
            row.cells[i].width=Cm(w); cant_split(row)


def clear(cell):
    p=cell.paragraphs[0]; p.clear(); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
    return p


def bottom_line(p):
    pr=p._p.get_or_add_pPr(); pb=OxmlElement('w:pBdr'); bt=OxmlElement('w:bottom')
    bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'5'); bt.set(qn('w:color'),BORDER); bt.set(qn('w:space'),'1')
    pb.append(bt); pr.append(pb)


def card(cell,title,desc='',opts=None,labels=None):
    margins(cell); border(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
    p=clear(cell); r=p.add_run(title); r.bold=True; r.font.name='Arial'; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(ACCENT)
    p.paragraph_format.space_after=Pt(1.2)
    if desc:
        p=cell.add_paragraph(); p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=1
        r=p.add_run(desc); r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(DARK)
    for opt in opts or []:
        p=cell.add_paragraph(); p.paragraph_format.space_after=Pt(.35); p.paragraph_format.line_spacing=1
        r=p.add_run(opt); r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(DARK)
    if labels:
        for label in labels:
            p=cell.add_paragraph(); p.paragraph_format.space_after=Pt(.3)
            r=p.add_run(label); r.bold=True; r.font.name='Arial'; r.font.size=Pt(7.6); r.font.color.rgb=RGBColor.from_string(MUTED)
            p=cell.add_paragraph(); p.paragraph_format.space_after=Pt(.8); bottom_line(p)
    elif opts is None:
        p=cell.add_paragraph(); p.paragraph_format.space_after=Pt(0); bottom_line(p)


def pidx(doc,prefix):
    for i,p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(prefix): return i
    raise RuntimeError('Missing '+prefix)


def grid_paragraph_section(doc,start,end,with_desc):
    ps=doc.paragraphs; a=pidx(doc,start); b=pidx(doc,end); instruction=ps[a+1]
    src=ps[a+2:b]; step=5 if with_desc else 4
    if len(src)%step: raise RuntimeError('Unexpected block count '+start)
    blocks=[]
    for i in range(0,len(src),step):
        if with_desc: blocks.append((src[i].text,src[i+1].text,[src[i+2].text,src[i+3].text,src[i+4].text]))
        else: blocks.append((src[i].text,'',[src[i+1].text,src[i+2].text,src[i+3].text]))
    t=doc.add_table(rows=(len(blocks)+1)//2,cols=2); fixed(t,[8.6,8.6])
    for i,bk in enumerate(blocks): card(t.cell(i//2,i%2),*bk)
    instruction._p.addnext(t._tbl)
    for p in src:
        if p._p.getparent() is not None: p._p.getparent().remove(p._p)


def between(doc,start,end):
    body=list(doc.element.body); si=ei=None
    for i,el in enumerate(body):
        if el.tag!=qn('w:p'): continue
        txt=Paragraph(el,doc).text.strip()
        if si is None and txt.startswith(start): si=i
        elif si is not None and txt.startswith(end): ei=i; break
    if si is None or ei is None: raise RuntimeError('Missing range '+start)
    return body[si+1:ei]


def compact_altgr(doc):
    elems=between(doc,'3. Sonderzeichen mit AltGr','4. Programme & Browser')
    instruction=next(Paragraph(e,doc) for e in elems if e.tag==qn('w:p') and Paragraph(e,doc).text.strip())
    old=next(Table(e,doc) for e in elems if e.tag==qn('w:tbl'))
    items=[row.cells[0].text.strip() for row in old.rows]
    t=doc.add_table(rows=4,cols=4); fixed(t,[2.0,6.6,2.0,6.6])
    for r,row in enumerate(t.rows):
        for cell in row.cells:
            margins(cell,45,70,45,70); border(cell,'E2E8F0'); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for side in range(2):
            lc=row.cells[side*2]; fc=row.cells[side*2+1]
            p=clear(lc); rr=p.add_run(items[r*2+side]); rr.bold=True; rr.font.name='Arial'; rr.font.size=Pt(8.5); rr.font.color.rgb=RGBColor.from_string(DARK)
            clear(fc); shade(fc)
    instruction._p.addnext(t._tbl); old._tbl.getparent().remove(old._tbl)


def grid_answer_tables(doc,start,end):
    elems=between(doc,start,end); instruction=next(Paragraph(e,doc) for e in elems if e.tag==qn('w:p') and Paragraph(e,doc).text.strip())
    old=[Table(e,doc) for e in elems if e.tag==qn('w:tbl')]
    labels=[t.cell(0,0).text.strip() for t in old]
    t=doc.add_table(rows=(len(labels)+1)//2,cols=2); fixed(t,[8.6,8.6])
    for i,label in enumerate(labels): card(t.cell(i//2,i%2),label,opts=None)
    instruction._p.addnext(t._tbl)
    for x in old: x._tbl.getparent().remove(x._tbl)


def grid_qa(doc,start,end,transfer=False):
    elems=between(doc,start,end); instruction=None; items=[]; cur=None
    for el in elems:
        if el.tag!=qn('w:p'): continue
        p=Paragraph(el,doc); txt=p.text.strip()
        if not txt: continue
        if instruction is None: instruction=p; continue
        if p.style.name=='Heading 2':
            if cur: items.append(cur)
            cur=[txt,'']
        elif cur is not None and not cur[1]: cur[1]=txt
    if cur: items.append(cur)
    t=doc.add_table(rows=(len(items)+1)//2,cols=2); fixed(t,[8.6,8.6])
    for i,(title,q) in enumerate(items): card(t.cell(i//2,i%2),title,q,labels=['Tastenkombination','Begründung'] if transfer else None,opts=None)
    instruction._p.addnext(t._tbl)
    for el in elems:
        if el is instruction._p: continue
        if el.getparent() is not None: el.getparent().remove(el)


def tighten(doc):
    for p in doc.paragraphs:
        txt=p.text.strip()
        if any(txt.startswith(x) for x in ('4. Programme & Browser','5. Windows & Arbeitsalltag','6. Kurz erklären','7. Nicht verwechseln','8. Transfer','Bonus - max. +2 P')):
            p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(1.5)
        if txt.startswith('5. Windows & Arbeitsalltag'): p.paragraph_format.page_break_before=False
        if txt.startswith(('Wähle die passendste','Schreibe die passende','Beschreibe kurz die Wirkung','Erkläre jeweils kurz','Nenne die sinnvollste','Genau unterscheiden:')):
            p.paragraph_format.space_after=Pt(1.5)
    for p in list(doc.paragraphs):
        txt=p.text.strip()
        if txt.startswith('Genau unterscheiden:'):
            p.text='Genau unterscheiden: je +1 P. · Bonuspunkte zusätzlich zu 41 P · Maximalnote 6.0'
            for r in p.runs: r.font.name='Arial'; r.font.size=Pt(7.3); r.font.color.rgb=RGBColor.from_string(MUTED)
        elif txt.startswith('Bonuspunkte zusätzlich'):
            parent=p._p.getparent()
            if parent is not None: parent.remove(p._p)
    for t in doc.tables:
        if len(t.rows)==2 and len(t.columns)==2 and 'Bonus 1' in t.cell(0,0).text and 'Bonus 2' in t.cell(1,0).text:
            for row in t.rows:
                row.height=Cm(.82); cant_split(row)
                for cell in row.cells:
                    margins(cell,35,75,35,75)
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(.5)
                        for r in p.runs: r.font.size=Pt(7.9 if r.bold else 7.2)
            break


doc=Document(PATH)
grid_paragraph_section(doc,'1. Was macht dieses Tastenkürzel?','2. Tastenkombinationen im Alltag',False)
grid_paragraph_section(doc,'2. Tastenkombinationen im Alltag','3. Sonderzeichen mit AltGr',True)
compact_altgr(doc)
grid_paragraph_section(doc,'4. Programme & Browser','5. Windows & Arbeitsalltag',True)
grid_answer_tables(doc,'5. Windows & Arbeitsalltag','6. Kurz erklären')
grid_answer_tables(doc,'6. Kurz erklären','7. Nicht verwechseln')
grid_qa(doc,'7. Nicht verwechseln','8. Transfer')
grid_qa(doc,'8. Transfer','Bonus - max. +2 P',True)
tighten(doc)
doc.save(PATH)
print(PATH)
