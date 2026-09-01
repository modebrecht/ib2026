from pathlib import Path
import sys

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

MIN_PT = 10.0
FORMULA_PREFIX = 'Linear: Note ='


def iter_paragraphs(container):
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def all_paragraphs(doc):
    # Do not deduplicate using Python object ids: python-docx creates temporary
    # wrapper objects and their ids can be reused while traversing, which can
    # accidentally skip unrelated paragraphs. Processing a repeated header or
    # cell twice is harmless; skipping visible text is not.
    yield from iter_paragraphs(doc)
    for section in doc.sections:
        for container in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            yield from iter_paragraphs(container)


def iter_cells(container):
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell
                yield from iter_cells(cell)


def all_cells(doc):
    # Duplicate processing is safe and avoids the same temporary-wrapper id
    # reuse problem as paragraph traversal.
    yield from iter_cells(doc)
    for section in doc.sections:
        for container in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            yield from iter_cells(container)


def normalize_cell_padding(cell, fallback=90):
    pr = cell._tc.get_or_add_tcPr()
    margins = pr.find(qn('w:tcMar'))
    if margins is None:
        margins = OxmlElement('w:tcMar')
        pr.append(margins)

    for logical_name, physical_name in [('start', 'left'), ('end', 'right')]:
        logical = margins.find(qn(f'w:{logical_name}'))
        physical = margins.find(qn(f'w:{physical_name}'))
        if physical is None:
            physical = OxmlElement(f'w:{physical_name}')
            margins.append(physical)

        value = logical.get(qn('w:w')) if logical is not None else None
        if not value:
            value = physical.get(qn('w:w'))
        if not value:
            value = str(fallback)

        physical.set(qn('w:w'), value)
        physical.set(qn('w:type'), 'dxa')


def effective_run_size_pt(doc, paragraph, run):
    if run.font.size is not None:
        return run.font.size.pt

    try:
        if run.style is not None and run.style.font.size is not None:
            return run.style.font.size.pt
    except Exception:
        pass

    try:
        if paragraph.style is not None and paragraph.style.font.size is not None:
            return paragraph.style.font.size.pt
    except Exception:
        pass

    normal = doc.styles['Normal'].font.size
    return normal.pt if normal is not None else None


def enforce(path):
    doc = Document(path)

    normal = doc.styles['Normal']
    if normal.font.size is None or normal.font.size.pt < MIN_PT:
        normal.font.size = Pt(MIN_PT)

    # Any explicitly smaller paragraph/character style is raised to the same
    # readability floor. The formula remains smaller via its explicit run size
    # and paragraph exception below.
    for style in doc.styles:
        try:
            if style.font.size is not None and style.font.size.pt < MIN_PT:
                style.font.size = Pt(MIN_PT)
        except Exception:
            pass

    for cell in all_cells(doc):
        normalize_cell_padding(cell)

    for paragraph in all_paragraphs(doc):
        if paragraph.text.strip().startswith(FORMULA_PREFIX):
            continue
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            size = effective_run_size_pt(doc, paragraph, run)
            if size is None or size < MIN_PT:
                run.font.size = Pt(MIN_PT)

    doc.save(path)

    # Hard validation: future generator/layout changes must not silently
    # reintroduce sub-10pt visible text anywhere except the formula line.
    check = Document(path)
    violations = []
    for paragraph in all_paragraphs(check):
        if paragraph.text.strip().startswith(FORMULA_PREFIX):
            continue
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            size = effective_run_size_pt(check, paragraph, run)
            if size is None or size + 1e-6 < MIN_PT:
                violations.append((paragraph.text.strip()[:80], run.text.strip()[:40], size))

    if violations:
        for paragraph_text, run_text, size in violations[:20]:
            print(f'FONT VIOLATION: {size} pt | {paragraph_text!r} | {run_text!r}', file=sys.stderr)
        raise SystemExit(f'{path}: {len(violations)} visible runs below {MIN_PT:g} pt')

    print(f'{path}: all visible text >= {MIN_PT:g} pt except formula; table padding normalized')


def main():
    if len(sys.argv) < 2:
        raise SystemExit('Usage: python scripts/enforce_p1_docx_readability.py <docx> [<docx> ...]')
    for value in sys.argv[1:]:
        path = Path(value)
        if not path.exists():
            raise SystemExit(f'{path} not found')
        enforce(path)


if __name__ == '__main__':
    main()
