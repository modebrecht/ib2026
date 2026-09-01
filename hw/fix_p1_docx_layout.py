from pathlib import Path
import re

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PATH = Path('hw/P1.docx')
SAFE_WIDTH_CM = 17.2


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:type'), 'dxa')
    tc_w.set(qn('w:w'), str(Cm(width_cm).twips))
    cell.width = Cm(width_cm)


def set_table_widths(table, widths_cm):
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tbl_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')

    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_w.set(qn('w:w'), str(Cm(sum(widths_cm)).twips))

    grid_cols = table._tbl.tblGrid.gridCol_lst
    for i, width in enumerate(widths_cm):
        table.columns[i].width = Cm(width)
        if i < len(grid_cols):
            grid_cols[i].set(qn('w:w'), str(Cm(width).twips))

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn('w:cantSplit')) is None:
            tr_pr.append(OxmlElement('w:cantSplit'))
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                set_cell_width(cell, widths_cm[i])


def normalize_cell_margins(cell, fallback_horizontal=100):
    """Materialize physical left/right cell padding for Word/LibreOffice.

    The base generator historically wrote logical start/end margins. Some
    renderers ignore those on table cells, which makes left-aligned text sit
    directly on the cell edge. Preserve the intended values by copying
    start/end to left/right; cells without explicit horizontal margins get a
    small, consistent fallback inset.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn('w:tcMar'))
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)

    for logical_name, physical_name in [('start', 'left'), ('end', 'right')]:
        logical = tc_mar.find(qn(f'w:{logical_name}'))
        physical = tc_mar.find(qn(f'w:{physical_name}'))
        if physical is None:
            physical = OxmlElement(f'w:{physical_name}')
            tc_mar.append(physical)

        width = None
        if logical is not None:
            width = logical.get(qn('w:w'))
        if not width:
            width = physical.get(qn('w:w'))
        if not width:
            width = str(fallback_horizontal)

        physical.set(qn('w:w'), width)
        physical.set(qn('w:type'), 'dxa')


def normalize_table_cell_margins(table):
    for row in table.rows:
        for cell in row.cells:
            normalize_cell_margins(cell)
            for nested_table in cell.tables:
                normalize_table_cell_margins(nested_table)


def set_paragraph_bottom_border(paragraph, color='94A3B8', size='6'):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn('w:pBdr'))
    if p_bdr is None:
        p_bdr = OxmlElement('w:pBdr')
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn('w:bottom'))
    if bottom is None:
        bottom = OxmlElement('w:bottom')
        p_bdr.append(bottom)
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)


def make_blank_writable_line(paragraph, in_cell=False):
    paragraph.clear()
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(4 if not in_cell else 1)
    paragraph.paragraph_format.line_spacing = 1.0
    set_paragraph_bottom_border(paragraph, color='CBD5E1' if in_cell else '94A3B8', size='5')


def walk_paragraphs(container, in_cell=False):
    for paragraph in container.paragraphs:
        yield paragraph, in_cell
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from walk_paragraphs(cell, True)


def table_width_profile(table):
    columns = len(table.columns)
    text = ' | '.join(cell.text for row in table.rows for cell in row.cells)

    if columns == 1:
        return [SAFE_WIDTH_CM]
    if columns == 3:
        return [7.9, 4.0, 5.3]
    if columns == 6:
        return [2.86] * 6
    if columns == 2:
        if 'Situation' in text and 'Antwort (E / V / A)' in text:
            return [13.7, 3.5]
        if 'Moderner Monitor oder Fernseher' in text or 'Interne SSD direkt auf dem Mainboard' in text:
            return [13.6, 3.6]
        return [8.6, 8.6]

    return [SAFE_WIDTH_CM / columns] * columns


if not PATH.exists():
    raise SystemExit(f'{PATH} not found')

doc = Document(PATH)

for table in doc.tables:
    set_table_widths(table, table_width_profile(table))
    normalize_table_cell_margins(table)

# Any paragraph that consists only of underscore placeholders becomes a real
# blank writable line. Students can click and type immediately; there is no
# placeholder text to select/delete first. This includes name fields, EVA
# answer cells, port answers and open-response lines.
for paragraph, in_cell in walk_paragraphs(doc):
    text = paragraph.text.strip()
    if re.fullmatch(r'_{4,}', text):
        make_blank_writable_line(paragraph, in_cell=in_cell)

# Remove all literal page breaks introduced by the base generator. Do not
# force Green IT onto a new page: letting Word reflow naturally keeps the last
# pages compact and avoids nearly empty pages after small content changes.
for paragraph in list(doc.paragraphs):
    page_breaks = paragraph._p.xpath('.//w:br[@w:type="page"]')
    if not page_breaks:
        continue
    for br in page_breaks:
        br.getparent().remove(br)
    if not paragraph.text.strip() and not paragraph._p.xpath('.//w:drawing'):
        paragraph._element.getparent().remove(paragraph._element)

# Task 5 should never be stranded at the bottom of the previous page. Start
# the complete RAM/HDD/SSD section on a fresh page while leaving the rest of
# the compact reflow untouched.
storage_heading = next((p for p in doc.paragraphs if p.text.strip().startswith('5. RAM, HDD oder SSD?')), None)
if storage_heading is not None:
    storage_heading.paragraph_format.page_break_before = True

green_heading = next((p for p in doc.paragraphs if p.text.strip().startswith('8. Green IT')), None)
if green_heading is not None:
    green_heading.paragraph_format.page_break_before = False

# Slightly tighten task spacing without changing the visual hierarchy.
for paragraph in doc.paragraphs:
    if paragraph.style and paragraph.style.name == 'Heading 1':
        paragraph.paragraph_format.space_before = Pt(7)
        paragraph.paragraph_format.space_after = Pt(4)
    elif paragraph.style and paragraph.style.name == 'Heading 2':
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(2)

doc.save(PATH)
print(PATH)