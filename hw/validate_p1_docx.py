from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn

PATH = Path('hw/P1.docx')
MAX_UNDERSCORES = 12
SAFETY_TWIPS = 120  # ~2.1 mm inside the document margins


def fail(message):
    raise SystemExit(f'P1 DOCX QA failed: {message}')


def all_paragraphs(container):
    for p in container.paragraphs:
        yield p
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from all_paragraphs(cell)


def usable_width_twips(section):
    return int((section.page_width - section.left_margin - section.right_margin) / 635)


def table_grid_width_twips(table):
    total = 0
    for grid_col in table._tbl.tblGrid.gridCol_lst:
        value = grid_col.get(qn('w:w'))
        if value:
            total += int(value)
    return total


if not PATH.exists():
    fail(f'{PATH} does not exist')

doc = Document(PATH)
if not doc.sections:
    fail('document has no section')

usable = min(usable_width_twips(section) for section in doc.sections)
for idx, table in enumerate(doc.tables, 1):
    width = table_grid_width_twips(table)
    if width <= 0:
        fail(f'table {idx} has no explicit grid width')
    if width > usable - SAFETY_TWIPS:
        fail(f'table {idx} grid is {width} twips; usable width is {usable} twips')
    layout = table._tbl.tblPr.find(qn('w:tblLayout'))
    if layout is None or layout.get(qn('w:type')) != 'fixed':
        fail(f'table {idx} is not fixed-layout')

for p in all_paragraphs(doc):
    if re.search(r'_{%d,}' % (MAX_UNDERSCORES + 1), p.text):
        fail(f'long unbreakable underscore sequence remains: {p.text[:50]!r}')

if any(p._p.xpath('.//w:br[@w:type="page"]') for p in doc.paragraphs):
    fail('literal page-break paragraph remains and may create a blank page after reflow')

green_heading = next((p for p in doc.paragraphs if p.text.strip().startswith('8. Green IT')), None)
if green_heading is None or not green_heading.paragraph_format.page_break_before:
    fail('Green IT page break is not attached to the heading')

if not any('Bonus 1 – Hardware Detective' in p.text for p in doc.paragraphs):
    fail('bonus section missing')

print(f'P1 DOCX QA passed: {len(doc.tables)} tables, fixed safe widths, no overflow-prone rule text.')
