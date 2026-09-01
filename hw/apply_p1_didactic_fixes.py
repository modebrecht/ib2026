from docx import Document
from docx.shared import Pt

PATH = 'hw/P1.docx'


def set_cell_text(cell, text, size=9.0):
    cell.text = text
    if cell.paragraphs and cell.paragraphs[0].runs:
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(size)


def replace_paragraph(doc, startswith, text, size=9.3):
    paragraph = next((p for p in doc.paragraphs if p.text.strip().startswith(startswith)), None)
    if paragraph is None:
        raise SystemExit(f'P1 didactic fix failed: paragraph not found: {startswith!r}')
    paragraph.text = text
    if paragraph.runs:
        paragraph.runs[0].font.name = 'Arial'
        paragraph.runs[0].font.size = Pt(size)
    return paragraph


doc = Document(PATH)

# 1) EVA: remove the answer-pattern cue E-V-A / E-V-A by mixing the six
# already-taught situations. The task still tests exactly the same concept.
eva_items = [
    'Der Preis erscheint auf dem Display.',
    'Du bewegst die Maus.',
    'Das Kassensystem berechnet den Gesamtpreis.',
    'Ein Barcode wird an der Kasse eingescannt.',
    'Der Mauszeiger erscheint an einer neuen Stelle auf dem Bildschirm.',
    'Der Computer berechnet die neue Position des Mauszeigers.',
]

eva_table = next((
    table for table in doc.tables
    if table.rows
    and len(table.rows[0].cells) == 2
    and table.rows[0].cells[0].text.strip() == 'Situation'
    and 'Antwort' in table.rows[0].cells[1].text
), None)
if eva_table is None or len(eva_table.rows) != 7:
    raise SystemExit('P1 didactic fix failed: EVA table not found or has unexpected shape')

for idx, text in enumerate(eva_items, 1):
    set_cell_text(eva_table.rows[idx].cells[0], f'{idx}. {text}', 9.2)

# 2) EVA verstehen: the old third point repeated the first point (monitor =
# output). Keep 41 total points, but make point 3 assess the missing input
# stage instead of rewarding the same insight twice. The answer goes on the
# blank line below, so students can type immediately without deleting blanks.
replace_paragraph(
    doc,
    '3. Was macht der Monitor mit dem bereits berechneten Ergebnis?',
    '3. Nenne ein passendes Beispiel für die Eingabe dieser EVA-Kette: Grafikkarte/CPU verarbeitet → Monitor zeigt das Ergebnis.',
)

# 3) Troubleshooting: test troubleshooting rather than negation and unknown
# specialist vocabulary. Students positively select exactly two sensible
# first steps from familiar hardware actions.
replace_paragraph(
    doc,
    'Bei jedem Fall gehören zwei Aktionen nicht zu den sinnvollen ersten Troubleshooting-Schritten.',
    'Kreuze bei jedem Fall genau zwei sinnvolle erste Troubleshooting-Schritte an. Je korrekt gewähltem Schritt 1 P. Werden in einem Fall mehr als zwei Schritte angekreuzt, gibt dieser Fall 0 P.',
    9.5,
)

troubleshooting_options = [
    [
        'Prüfen, ob das Stromkabel fest am PC und an der Steckdose steckt',
        'Eine andere Steckdose testen',
        'HDMI-/DisplayPort-Kabel zum Monitor prüfen',
        'Maus an einem anderen USB-Anschluss testen',
        'Netzwerkkabel austauschen',
        'Browser neu öffnen',
    ],
    [
        'HDMI-/DisplayPort-Kabel prüfen',
        'Richtigen Eingang am Monitor auswählen',
        'Netzwerkkabel zum Router prüfen',
        'Maus neu anschliessen',
        'Drucker aus- und einschalten',
        'Eine andere Webseite öffnen',
    ],
    [
        'Prüfen, ob das RJ45-Kabel am PC und am Router/Switch richtig steckt',
        'Ein anderes Netzwerkkabel testen',
        'HDMI-/DisplayPort-Kabel zum Monitor prüfen',
        'WLAN-Passwort neu eingeben',
        'Lautsprecher neu anschliessen',
        'Bildschirmhelligkeit ändern',
    ],
]

trouble_tables = []
for table in doc.tables:
    if len(table.rows) != 3 or any(len(row.cells) != 2 for row in table.rows):
        continue
    texts = [cell.text.strip() for row in table.rows for cell in row.cells]
    if len(texts) == 6 and all(text.startswith(('□', '☐')) for text in texts):
        trouble_tables.append(table)

if len(trouble_tables) != 3:
    raise SystemExit(f'P1 didactic fix failed: expected 3 troubleshooting tables, found {len(trouble_tables)}')

for table, options in zip(trouble_tables, troubleshooting_options):
    cells = [cell for row in table.rows for cell in row.cells]
    for cell, option in zip(cells, options):
        set_cell_text(cell, '□  ' + option, 8.8)

doc.save(PATH)
print(PATH)
